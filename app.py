import io, os, json, time
from copy import deepcopy

import streamlit as st
import requests
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

# -------------------- Config --------------------
st.set_page_config(page_title="One-Minute Resume Tailor", page_icon="📝", layout="centered")
st.title("📝 One-Minute Resume Tailor (Hugging Face)")
st.caption("Upload .docx resume + paste JD → get edited .docx that keeps your template. Export to PDF in Word/Google Docs.")

HF_TOKEN = st.secrets.get("HF_TOKEN") or os.environ.get("HF_TOKEN", "")
MODEL_ID = "HuggingFaceH4/zephyr-7b-beta"
API_URL = f"https://api-inference.huggingface.co/models/{MODEL_ID}"
HEADERS = {"Authorization": f"Bearer {HF_TOKEN}"}

if not HF_TOKEN:
    st.warning('⚠️ Add your HF token in Secrets: HF_TOKEN = "hf_..."')

# -------------------- Helpers --------------------
def _iter_block_items(parent):
    """Yield paragraphs and tables in document order (robust)."""
    try:
        parent_elm = parent.element.body
    except AttributeError:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_text_snapshot(doc):
    """Plain-text snapshot for the LLM, preserving section order."""
    lines = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = block.text.strip()
            if txt:
                lines.append(txt)
        else:  # table
            for row in block.rows:
                row_txt = " | ".join([cell.text.strip() for cell in row.cells])
                if row_txt.strip():
                    lines.append(row_txt)
    return "\n".join(lines)

SYSTEM_PROMPT = """You tailor resumes to a specific job description while preserving structure and tone.
Rules:
- Keep the resume ONE PAGE in content length.
- Preserve section order and headings from the original (Summary, Experience, Projects, Education, Skills).
- Rewrite bullets to mirror JD responsibilities and keywords naturally (no keyword stuffing).
- Prefer concise, metric-led bullets: Action Verb + What + Impact (+ Tools).
- Do NOT invent employers, titles, or dates; tighten and rephrase only.
Return only the revised resume text; no commentary.
"""

def call_hf_llm(resume_text: str, jd_text: str, temperature: float = 0.2, max_new_tokens: int = 900) -> str:
    """
    Calls Hugging Face Inference API for text generation with chat-style prompt.
    Handles cold starts (503) and common errors gracefully.
    """
    if not HF_TOKEN:
        raise RuntimeError("Missing HF token.")

    user_prompt = f"""
ORIGINAL RESUME (text-only snapshot):
---
{resume_text}
---

JOB DESCRIPTION:
---
{jd_text}
---

TASK:
Return ONLY the revised resume TEXT (no extra commentary), keeping the SAME SECTION ORDER and roughly the same number of bullets per experience. Keep it to ONE PAGE worth of concise content.
"""

    # Zephyr supports chat template via 'inputs' with 'messages'
    payload = {
        "inputs": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt}
        ],
        "parameters": {
            "max_new_tokens": max_new_tokens,
            "temperature": temperature,
            "top_p": 0.95,
            "repetition_penalty": 1.05,
            "return_full_text": False
        },
        "options": {
            "use_cache": True,
            "wait_for_model": True  # handle cold starts on free tier
        }
    }

    # Basic retry loop for cold starts / transient 5xx
    for attempt in range(5):
        resp = requests.post(API_URL, headers=HEADERS, json=payload, timeout=120)
        if resp.status_code == 200:
            try:
                data = resp.json()
                # HF Inference returns a list of generated texts or a dict with 'generated_text'
                if isinstance(data, list) and data and "generated_text" in data[0]:
                    return data[0]["generated_text"].strip()
                # Some models return 'conversation' style objects
                if isinstance(data, dict) and "generated_text" in data:
                    return data["generated_text"].strip()
                # Fallback: try to decode text field
                text = data if isinstance(data, str) else json.dumps(data)
                return text.strip()
            except Exception:
                return resp.text.strip()

        # 503 = model loading; 429 = rate limit
        if resp.status_code in (503, 429):
            time.sleep(3 + attempt * 2)
            continue

        # other errors: raise with message
        try:
            err = resp.json()
        except Exception:
            err = {"error": resp.text}
        raise RuntimeError(f"HF API error {resp.status_code}: {err}")

    raise RuntimeError("HF API retry limit reached. Try again in a minute.")

def rewrite_doc_in_place(src_doc, revised_text):
    """Replace text while preserving styles as much as possible."""
    new_lines = [ln.rstrip() for ln in revised_text.splitlines()]
    it = iter(new_lines)

    def next_line():
        try:
            return next(it)
        except StopIteration:
            return ""

    for block in _iter_block_items(src_doc):
        if isinstance(block, Paragraph):
            line = next_line()
            if block.runs:
                block.runs[0].text = line
                for r in block.runs[1:]:
                    r.text = ""
            else:
                block.text = line
        else:  # table
            for row in block.rows:
                for cell in row.cells:
                    line = next_line()
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                        if p.runs:
                            p.runs[0].text = line
                            for r in p.runs[1:]:
                                r.text = ""
                        else:
                            p.text = line
    return src_doc

# -------------------- UI --------------------
resume_file = st.file_uploader("📄 Upload your **.docx** resume", type=["docx"])
jd = st.text_area("📝 Paste the Job Description", height=220, placeholder="Paste JD here...")

with st.expander("Advanced"):
    st.write("Model:", MODEL_ID)
    temp = st.slider("Creativity (temperature)", 0.0, 1.0, 0.2, 0.05)
    max_tokens = st.slider("Max new tokens", 200, 1200, 900, 50)

if st.button("✨ Tailor my resume", type="primary", use_container_width=True):
    if not (resume_file and jd and HF_TOKEN):
        st.error("❌ Please upload a .docx, paste a JD, and set HF_TOKEN in Secrets.")
        st.stop()

    try:
        src = DocxDocument(resume_file)
    except Exception as e:
        st.error(f"Could not read .docx. Ensure it is a valid Word file. Error: {e}")
        st.stop()

    snapshot = extract_text_snapshot(src)

    with st.spinner("⏳ Rewriting bullets and aligning to JD (Hugging Face)..."):
        try:
            revised_text = call_hf_llm(snapshot, jd, temperature=temp, max_new_tokens=max_tokens)
        except Exception as e:
            st.error(f"HF API error: {e}")
            st.stop()

    revised_doc = rewrite_doc_in_place(deepcopy(src), revised_text)

    buf = io.BytesIO()
    revised_doc.save(buf)

    st.success("✅ Done! Download your tailored resume:")
    st.download_button(
        "⬇️ Download .docx",
        data=buf.getvalue(),
        file_name="Tailored_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    st.info("💡 Tip: Open the .docx in Google Docs or Word → File → Download → PDF for a perfect export.")

# Optional quick token test in the sidebar
if st.sidebar.button("🔑 Test HF token"):
    if HF_TOKEN:
        st.sidebar.success("HF token loaded ✅")
    else:
        st.sidebar.error("HF token missing ❌")
